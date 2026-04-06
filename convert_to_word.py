#!/usr/bin/env python3
"""
Convert SYSTEM_DOCUMENTATION.md to Word document (.docx)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def create_word_document():
    # Create document
    doc = Document()
    
    # Set document styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Alafia Point Wellness Clinic - Management System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_heading('Comprehensive User Documentation', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        "1. System Overview",
        "2. System Modules", 
        "3. Getting Started",
        "4. Module Descriptions",
        "5. Key Features",
        "6. User Roles & Permissions",
        "7. Reports & Analytics",
        "8. System Customization",
        "9. Troubleshooting",
        "10. Support",
        "11. Quick Reference Guide"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. System Overview
    doc.add_heading('1. System Overview', 1)
    
    p = doc.add_paragraph()
    p.add_run('Alafia Point Wellness Clinic Management System').bold = True
    p.add_run(' is a comprehensive, web-based healthcare management solution designed to streamline all aspects of clinic operations. Built on modern Django framework with a responsive Bootstrap interface, the system provides an intuitive, professional experience for healthcare providers, administrators, and staff.')
    
    doc.add_heading('System Capabilities', 2)
    capabilities = [
        'Patient Management: Complete patient lifecycle from registration to discharge',
        'Appointment Scheduling: Calendar-based appointment management with notifications',
        'Medical Records: Digital health records with document upload capabilities',
        'Laboratory Management: Test ordering, processing, and reporting',
        'Billing & Invoicing: Comprehensive billing with insurance support and group invoicing',
        'Inventory Management: Drug and medical supplies tracking',
        'Staff Management: Employee records, roles, and permissions',
        'Financial Reporting: Revenue, expenses, and performance analytics'
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_page_break()
    
    # 2. System Modules
    doc.add_heading('2. System Modules', 1)
    
    p = doc.add_paragraph('The system consists of ')
    p.add_run('9 integrated modules').bold = True
    p.add_run(' that work together seamlessly:')
    
    # Create table for modules
    table = doc.add_table(rows=10, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Module'
    hdr_cells[1].text = 'Icon'
    hdr_cells[2].text = 'Description'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    modules_data = [
        ('Patients Management', '👥', 'Patient registration, profiles, group management'),
        ('Appointments', '📅', 'Scheduling, calendar views, reminders'),
        ('Medical Records', '📁', 'Digital health records, prescriptions, diagnoses'),
        ('Laboratory', '🧪', 'Lab test requests, results, certificates'),
        ('Pharmacy', '💊', 'Medication management, prescriptions'),
        ('Billing', '💳', 'Invoicing, payments, group billing, insurance'),
        ('Budget & Expenses', '💰', 'Financial tracking, expense management'),
        ('Staff Management', '👤', 'Employee records, departments, roles'),
        ('Reports', '📊', 'Analytics, statistics, data exports')
    ]
    
    for i, (module, icon, desc) in enumerate(modules_data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = module
        row_cells[1].text = icon
        row_cells[2].text = desc
    
    doc.add_page_break()
    
    # 3. Getting Started
    doc.add_heading('3. Getting Started', 1)
    
    doc.add_heading('System Access', 2)
    access_steps = [
        'Login: Access the system through your web browser at the provided URL',
        'Credentials: Use your assigned username and password',
        'Dashboard: Upon login, you\'ll see the main dashboard with quick access to all modules'
    ]
    for step in access_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('Navigation', 2)
    nav_items = [
        'Sidebar Menu: Access all modules from the left sidebar',
        'Quick Actions: Common tasks available from the dashboard',
        'Search: Global search functionality across patients, records, and documents',
        'User Profile: Access settings and logout from the top-right corner'
    ]
    for item in nav_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 4. Module Descriptions
    doc.add_heading('4. Module Descriptions', 1)
    
    # 4.1 Patients Management
    doc.add_heading('4.1 Patients Management', 2)
    
    doc.add_heading('Patient Registration', 3)
    reg_items = [
        'Register new patients with comprehensive demographic information',
        'Generate unique Patient IDs automatically',
        'Capture contact details, emergency contacts, and insurance information',
        'Upload patient photos and documents'
    ]
    for item in reg_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Patient Groups', 3)
    group_items = [
        'Create patient groups for differential pricing (e.g., Corporate, Insurance, VIP)',
        'Manage group-specific pricing for services and lab tests',
        'Assign patients to multiple groups',
        'View group statistics and billing summaries'
    ]
    for item in group_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Patient Dashboard', 3)
    dash_items = [
        'Complete patient history view',
        'Quick access to appointments, medical records, lab tests',
        'Billing history and payment status',
        'Allergies and medical alerts'
    ]
    for item in dash_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Key Features', 3)
    key_features = [
        'Advanced Search: Search by name, ID, phone, or email',
        'Status Tracking: Active, inactive, or discharged patients',
        'Export: Export patient data to Excel',
        'Group Management: Batch operations for patient groups'
    ]
    for item in key_features:
        doc.add_paragraph(item, style='List Bullet')
    
    # 4.2 Appointments
    doc.add_heading('4.2 Appointments', 2)
    
    doc.add_heading('Scheduling', 3)
    sched_items = [
        'Create appointments with date, time, and doctor selection',
        'Color-coded appointment status (Pending, Confirmed, Completed, Cancelled)',
        'Drag-and-drop calendar interface',
        'Recurring appointments support'
    ]
    for item in sched_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Calendar Views', 3)
    cal_items = [
        'Daily, weekly, and monthly calendar views',
        'Doctor-specific appointment filters',
        'Department-based scheduling',
        'Waiting room management'
    ]
    for item in cal_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Notifications', 3)
    notif_items = [
        'Automatic appointment reminders',
        'SMS and email notifications (configurable)',
        'Check-in/check-out tracking'
    ]
    for item in notif_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 4.3 Medical Records
    doc.add_heading('4.3 Medical Records', 2)
    
    doc.add_heading('Record Types', 3)
    record_types = [
        'Diagnosis: ICD-coded diagnoses with descriptions',
        'Treatment Plans: Detailed treatment protocols',
        'Prescriptions: Medication orders with dosage instructions',
        'Lab Results: Integrated laboratory test results',
        'Clinical Notes: Free-text clinical observations',
        'Vital Signs: Temperature, BP, pulse, etc.'
    ]
    for item in record_types:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Document Management', 3)
    doc_items = [
        'Upload medical images, X-rays, scans',
        'Attach external documents',
        'Organize by record type and date',
        'Secure document storage'
    ]
    for item in doc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Features', 3)
    feat_items = [
        'Timeline View: Chronological medical history',
        'Print Records: Generate printable medical reports',
        'Export: Export to PDF or Excel',
        'Access Control: Role-based record access'
    ]
    for item in feat_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 4.4 Laboratory
    doc.add_heading('4.4 Laboratory', 2)
    
    doc.add_heading('Test Management', 3)
    test_items = [
        'Test Catalog: Comprehensive list of available tests with pricing',
        'Test Requests: Order tests for patients with priority levels',
        'Sample Collection: Track sample collection status',
        'Results Entry: Enter test results with reference ranges',
        'Report Generation: Generate professional test reports'
    ]
    for item in test_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Lab Dashboard', 3)
    lab_dash = [
        'Pending tests queue',
        'Completed tests ready for review',
        'Overdue tests alerts',
        'Daily workload statistics'
    ]
    for item in lab_dash:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Special Features', 3)
    spec_items = [
        'Batch Processing: Process multiple tests simultaneously',
        'Certificates: Generate lab certificates for patients',
        'Report Templates: Customizable report formats',
        'Email Reports: Send reports directly to patients'
    ]
    for item in spec_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 4.5 Pharmacy
    doc.add_heading('4.5 Pharmacy', 2)
    
    doc.add_heading('Medication Management', 3)
    med_items = [
        'Drug catalog with comprehensive information',
        'Stock level tracking',
        'Expiry date monitoring',
        'Batch number tracking'
    ]
    for item in med_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Prescriptions', 3)
    presc_items = [
        'Create prescriptions from medical records',
        'Dosage and frequency instructions',
        'Drug interaction warnings',
        'Print prescription slips'
    ]
    for item in presc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Dispensing', 3)
    disp_items = [
        'Track dispensed medications',
        'Patient medication history',
        'Refill management',
        'Drug usage reports'
    ]
    for item in disp_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 4.6 Billing
    doc.add_heading('4.6 Billing', 2)
    
    doc.add_heading('Invoice Management', 3)
    inv_items = [
        'Individual Invoices: Create invoices for services and treatments',
        'Group Invoices: Bulk billing for corporate/insurance groups',
        'Invoice Status: Draft, Sent, Paid, Overdue, Cancelled',
        'Payment Tracking: Record payments against invoices'
    ]
    for item in inv_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Payment Processing', 3)
    pay_items = [
        'Multiple payment methods (Cash, Card, Mobile Money, Insurance)',
        'Partial payment support',
        'Overpayment handling',
        'Refund processing'
    ]
    for item in pay_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Group Billing', 3)
    gb_items = [
        'Generate periodic invoices for patient groups',
        'Automatic invoice calculation based on services',
        'Group payment tracking',
        'Corporate billing summaries'
    ]
    for item in gb_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Financial Features', 3)
    fin_items = [
        'Aging Reports: Track overdue payments',
        'Revenue Reports: Daily, monthly, yearly revenue analysis',
        'Excel Export: Export billing data for accounting',
        'Payment Receipts: Generate payment receipts'
    ]
    for item in fin_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 4.7 Budget & Expenses
    doc.add_heading('4.7 Budget & Expenses', 2)
    
    doc.add_heading('Expense Tracking', 3)
    exp_items = [
        'Record operational expenses',
        'Categorize expenses (Utilities, Supplies, Salaries, etc.)',
        'Attach receipts and documents',
        'Expense approval workflows'
    ]
    for item in exp_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Budget Management', 3)
    bud_items = [
        'Set department budgets',
        'Track budget vs. actual spending',
        'Variance analysis',
        'Financial forecasting'
    ]
    for item in bud_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Cash Flow', 3)
    cf_items = [
        'Daily cash flow tracking',
        'Income and expense summaries',
        'Bank reconciliation support'
    ]
    for item in cf_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 4.8 Staff Management
    doc.add_heading('4.8 Staff Management', 2)
    
    doc.add_heading('Employee Records', 3)
    emp_items = [
        'Complete staff profiles',
        'Role-based access control',
        'Department assignments',
        'Employment history'
    ]
    for item in emp_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Attendance & Leave', 3)
    att_items = [
        'Attendance tracking',
        'Leave management',
        'Shift scheduling',
        'Payroll integration support'
    ]
    for item in att_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Performance', 3)
    perf_items = [
        'Staff performance metrics',
        'Patient feedback tracking',
        'Service quality monitoring'
    ]
    for item in perf_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Security', 3)
    sec_items = [
        'Role-based permissions',
        'User activity logs',
        'Password management',
        'Account security settings'
    ]
    for item in sec_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 4.9 Reports
    doc.add_heading('4.9 Reports', 2)
    
    doc.add_heading('Operational Reports', 3)
    op_items = [
        'Patient statistics',
        'Appointment analytics',
        'Laboratory workload',
        'Pharmacy usage'
    ]
    for item in op_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Financial Reports', 3)
    finr_items = [
        'Revenue analysis',
        'Expense reports',
        'Profit & Loss statements',
        'Cash flow reports'
    ]
    for item in finr_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Custom Reports', 3)
    cr_items = [
        'Date range filtering',
        'Department-wise reports',
        'Export to Excel/PDF',
        'Scheduled report generation'
    ]
    for item in cr_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 5. Key Features
    doc.add_heading('5. Key Features', 1)
    
    doc.add_heading('Unified Interface', 2)
    ui_items = [
        'Consistent Design: All tables use standardized styling with gradient headers, sortable columns, and hover effects',
        'Responsive Design: Works on desktop, tablet, and mobile devices',
        'Intuitive Navigation: Easy access to all features through logical menu structure'
    ]
    for item in ui_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Advanced Table Features', 2)
    table_items = [
        'Sorting: Click column headers to sort data',
        'Pagination: Navigate through large datasets with Next/Previous controls',
        'Filtering: Filter data by multiple criteria',
        'Export: Export table data to Excel or PDF',
        'Search: Quick search within tables'
    ]
    for item in table_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Group Management', 2)
    gm_items = [
        'Patient Groups: Organize patients into groups (Corporate, Insurance, etc.)',
        'Differential Pricing: Set custom prices for each group',
        'Group Invoicing: Generate consolidated invoices for groups',
        'Batch Operations: Apply operations to entire groups'
    ]
    for item in gm_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Document Management', 2)
    dm_items = [
        'Upload & Store: Securely store medical documents',
        'Organize: Categorize documents by type and patient',
        'Access: Role-based document access',
        'Download: Easy document retrieval'
    ]
    for item in dm_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Security & Compliance', 2)
    sc_items = [
        'Role-Based Access: Different permissions for different user roles',
        'Audit Trail: Track all system activities',
        'Data Backup: Automated database backups',
        'Secure Login: Encrypted password storage'
    ]
    for item in sc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 6. User Roles & Permissions
    doc.add_heading('6. User Roles & Permissions', 1)
    
    doc.add_heading('Default Roles', 2)
    
    # Create roles table
    roles_table = doc.add_table(rows=10, cols=3)
    roles_table.style = 'Light Grid Accent 1'
    
    # Header
    hdr = roles_table.rows[0].cells
    hdr[0].text = 'Role'
    hdr[1].text = 'Description'
    hdr[2].text = 'Access Level'
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    
    roles_data = [
        ('Superuser', 'System Administrator', 'Full system access'),
        ('Admin', 'Clinic Manager', 'All modules, user management'),
        ('Doctor', 'Medical Doctor', 'Patients, Appointments, Medical Records, Lab'),
        ('Nurse', 'Nursing Staff', 'Patients, Appointments, basic Medical Records'),
        ('Lab Technician', 'Laboratory Staff', 'Laboratory module, test results'),
        ('Receptionist', 'Front Desk', 'Patients, Appointments, basic Billing'),
        ('Pharmacist', 'Pharmacy Staff', 'Pharmacy, Inventory, prescriptions'),
        ('Accountant', 'Finance Staff', 'Billing, Budget, Reports'),
        ('Billing Staff', 'Billing Department', 'Billing, Payments, Invoicing')
    ]
    
    for i, (role, desc, access) in enumerate(roles_data, 1):
        row = roles_table.rows[i].cells
        row[0].text = role
        row[1].text = desc
        row[2].text = access
    
    doc.add_heading('Permission Levels', 2)
    perms = [
        'View: Can view records',
        'Create: Can add new records',
        'Edit: Can modify existing records',
        'Delete: Can remove records',
        'Export: Can export data',
        'Admin: Can manage system settings'
    ]
    for p in perms:
        doc.add_paragraph(p, style='List Bullet')
    
    doc.add_page_break()
    
    # 7. Reports & Analytics
    doc.add_heading('7. Reports & Analytics', 1)
    
    doc.add_heading('Available Reports', 2)
    
    doc.add_heading('Patient Reports', 3)
    pr_items = [
        'Patient registration statistics',
        'Patient demographics',
        'Patient group analysis',
        'Patient visit history'
    ]
    for item in pr_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Appointment Reports', 3)
    ar_items = [
        'Daily/weekly appointment schedules',
        'Doctor utilization',
        'No-show analysis',
        'Waiting time reports'
    ]
    for item in ar_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Financial Reports', 3)
    fr_items = [
        'Daily revenue summary',
        'Monthly billing report',
        'Outstanding payments (Aging report)',
        'Group billing summary',
        'Expense analysis'
    ]
    for item in fr_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Laboratory Reports', 3)
    lr_items = [
        'Test volume by type',
        'Turnaround time analysis',
        'Revenue by test category',
        'Pending tests report'
    ]
    for item in lr_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Custom Reports', 3)
    cus_items = [
        'Build custom queries',
        'Save report templates',
        'Schedule automated reports',
        'Export in multiple formats'
    ]
    for item in cus_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 8. System Customization
    doc.add_heading('8. System Customization', 1)
    
    doc.add_heading('Clinic Settings', 2)
    p = doc.add_paragraph('Access ')
    p.add_run('Settings > Clinic Settings').bold = True
    p.add_run(' to customize:')
    
    doc.add_heading('Branding', 3)
    brand_items = [
        'Clinic Name: Change displayed clinic name',
        'Logo: Upload clinic logo (recommended: 200x80px)',
        'Contact Info: Address, phone, email, website'
    ]
    for item in brand_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Theme Customization', 3)
    p = doc.add_paragraph('Customize colors to match your brand:')
    theme_items = [
        'Primary Color: Main brand color',
        'Success Color: For positive actions',
        'Warning Color: For cautions',
        'Danger Color: For errors/deletions',
        'Info Color: For informational messages'
    ]
    for item in theme_items:
        doc.add_paragraph(item, style='List Bullet')
    p = doc.add_paragraph()
    p.add_run('All colors support hex codes (e.g., #1B5E96).').italic = True
    
    doc.add_heading('Module Management', 3)
    mm_items = [
        'Toggle modules on/off',
        'Reorder modules in navigation',
        'Customize module display names'
    ]
    for item in mm_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Data Management', 2)
    
    doc.add_heading('Backup', 3)
    p = doc.add_paragraph()
    p.add_run('The system provides robust backup capabilities to ensure your data is always protected:').bold = True
    
    doc.add_heading('Automated Backups', 4)
    auto_items = [
        'Daily Backups: Automatic daily database backups at 2:00 AM',
        'Retention Policy: Maintains last 30 days of backups',
        'Storage Location: Backups stored in db_backups/ folder',
        'File Format: SQLite database files (.db format)'
    ]
    for item in auto_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Manual Backup', 4)
    p = doc.add_paragraph('To create a manual backup at any time:')
    manual_steps = [
        'Navigate to Settings > System Maintenance',
        'Click "Create Backup Now" button',
        'The backup will be created immediately with timestamp',
        'Download the backup file for off-site storage'
    ]
    for step in manual_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('Backup Schedule', 4)
    # Backup schedule table
    bs_table = doc.add_table(rows=4, cols=3)
    bs_table.style = 'Light Grid Accent 1'
    
    bs_hdr = bs_table.rows[0].cells
    bs_hdr[0].text = 'Type'
    bs_hdr[1].text = 'Frequency'
    bs_hdr[2].text = 'Retention'
    for cell in bs_hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    
    bs_data = [
        ('Daily Auto', 'Every day at 2:00 AM', '30 days'),
        ('Manual', 'On-demand', 'Permanent'),
        ('Pre-Update', 'Before system updates', 'Until next update')
    ]
    
    for i, (typ, freq, ret) in enumerate(bs_data, 1):
        row = bs_table.rows[i].cells
        row[0].text = typ
        row[1].text = freq
        row[2].text = ret
    
    doc.add_heading('Backup Contents', 4)
    bc_items = [
        'Patient records and demographics',
        'Medical records and documents',
        'Appointments and schedules',
        'Billing and payment data',
        'Laboratory results',
        'Staff information',
        'System settings and configurations',
        'User accounts and permissions'
    ]
    for item in bc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Note: ').bold = True
    p.add_run('Media files (uploaded documents, images) are stored separately and should be backed up via file system backup.')
    
    doc.add_heading('Restoring from Backup', 4)
    
    p = doc.add_paragraph()
    p.add_run('Via Admin Interface:').bold = True
    restore_steps = [
        'Go to Settings > System Maintenance > Restore',
        'Select the backup file from the list',
        'Click "Restore Database"',
        'Confirm the restoration (WARNING: Current data will be replaced)',
        'System will restart after restoration'
    ]
    for step in restore_steps:
        doc.add_paragraph(step, style='List Number')
    
    p = doc.add_paragraph()
    p.add_run('Via Command Line:').bold = True
    p = doc.add_paragraph()
    p.add_run('# Stop the server first\npython manage.py backup_db --restore --file=backup_filename.db').font.name = 'Courier New'
    
    doc.add_heading('Important Notes for Restoration', 4)
    rest_notes = [
        'Downtime Required: System will be unavailable during restore',
        'Data Loss Risk: Current data after backup date will be lost',
        'Verification: Always verify restored data integrity',
        'Backup First: Create fresh backup before any restoration'
    ]
    for item in rest_notes:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Backup Best Practices', 4)
    bp_items = [
        'Regular Testing: Test restore process monthly',
        'Off-site Storage: Download backups to external storage weekly',
        'Multiple Copies: Maintain 3 copies (local, external, cloud)',
        'Documentation: Record backup dates and contents',
        'Verification: Verify backup file integrity after creation'
    ]
    for item in bp_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Backup Monitoring', 4)
    bm_items = [
        'View backup status on the Dashboard',
        'Receive email alerts for backup failures',
        'Check backup logs in Settings > Logs'
    ]
    for item in bm_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Import/Export', 3)
    ie_items = [
        'Import patient data from Excel',
        'Export reports to Excel/PDF',
        'Bulk data operations'
    ]
    for item in ie_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 9. Troubleshooting
    doc.add_heading('9. Troubleshooting', 1)
    
    doc.add_heading('Common Issues & Solutions', 2)
    
    doc.add_heading('Cannot Login', 3)
    login_issues = [
        'Check credentials: Verify username and password',
        'Caps Lock: Ensure Caps Lock is off',
        'Account Status: Contact admin to verify account is active'
    ]
    for item in login_issues:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Page Not Loading', 3)
    page_issues = [
        'Clear Cache: Clear browser cache and cookies',
        'Browser: Use modern browsers (Chrome, Firefox, Edge)',
        'Internet: Check internet connection'
    ]
    for item in page_issues:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Data Not Saving', 3)
    data_issues = [
        'Required Fields: Ensure all required fields are filled',
        'Valid Data: Check data format (dates, numbers, etc.)',
        'Permissions: Verify you have edit permissions'
    ]
    for item in data_issues:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Reports Not Generating', 3)
    report_issues = [
        'Date Range: Ensure valid date range selected',
        'Data Available: Verify data exists for selected period',
        'Export Format: Check selected export format'
    ]
    for item in report_issues:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Error Messages', 2)
    
    # Error messages table
    err_table = doc.add_table(rows=5, cols=2)
    err_table.style = 'Light Grid Accent 1'
    
    err_hdr = err_table.rows[0].cells
    err_hdr[0].text = 'Error'
    err_hdr[1].text = 'Solution'
    for cell in err_hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    
    err_data = [
        ('"Permission Denied"', 'Contact administrator for access'),
        ('"Record Not Found"', 'Check ID or search with different criteria'),
        ('"Validation Error"', 'Review form fields for errors'),
        ('"Server Error"', 'Try again later or contact support')
    ]
    
    for i, (err, sol) in enumerate(err_data, 1):
        row = err_table.rows[i].cells
        row[0].text = err
        row[1].text = sol
    
    doc.add_page_break()
    
    # 10. Support
    doc.add_heading('10. Support', 1)
    
    doc.add_heading('Getting Help', 2)
    
    doc.add_heading('In-System Help', 3)
    ish_items = [
        'Tooltips: Hover over icons for quick tips',
        'Help Icons: Click help icons for detailed information',
        'Contextual Help: Module-specific guidance'
    ]
    for item in ish_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Technical Support', 3)
    p = doc.add_paragraph()
    p.add_run('Email: ').bold = True
    p.add_run('support@alafiapoint.com')
    p = doc.add_paragraph()
    p.add_run('Phone: ').bold = True
    p.add_run('+256-XXX-XXXXXX')
    p = doc.add_paragraph()
    p.add_run('Hours: ').bold = True
    p.add_run('Monday-Friday, 8:00 AM - 6:00 PM')
    
    doc.add_heading('Training Resources', 3)
    tr_items = [
        'Video Tutorials: Available on dashboard',
        'User Manual: This document',
        'FAQ Section: Common questions answered'
    ]
    for item in tr_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Feedback', 2)
    p = doc.add_paragraph('We value your feedback! Please submit suggestions and feature requests through:')
    fb_items = [
        'Feedback Form: Available in system menu',
        'Email: feedback@alafiapoint.com'
    ]
    for item in fb_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 11. Quick Reference Guide
    doc.add_heading('11. Quick Reference Guide', 1)
    
    doc.add_heading('Keyboard Shortcuts', 2)
    
    # Keyboard shortcuts table
    ks_table = doc.add_table(rows=5, cols=2)
    ks_table.style = 'Light Grid Accent 1'
    
    ks_hdr = ks_table.rows[0].cells
    ks_hdr[0].text = 'Shortcut'
    ks_hdr[1].text = 'Action'
    for cell in ks_hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    
    ks_data = [
        ('Ctrl + S', 'Save current form'),
        ('Ctrl + P', 'Print current page'),
        ('Ctrl + F', 'Find/Search on page'),
        ('Esc', 'Close modal/dialog')
    ]
    
    for i, (shortcut, action) in enumerate(ks_data, 1):
        row = ks_table.rows[i].cells
        row[0].text = shortcut
        row[1].text = action
    
    doc.add_heading('Quick Actions', 2)
    p = doc.add_paragraph('From the Dashboard:')
    qa_items = [
        '+ New Patient: Quick patient registration',
        '+ New Appointment: Schedule appointment',
        '+ New Invoice: Create invoice',
        'Search: Global search bar'
    ]
    for item in qa_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Status Colors', 2)
    
    # Status colors table
    sc_table = doc.add_table(rows=6, cols=2)
    sc_table.style = 'Light Grid Accent 1'
    
    sc_hdr = sc_table.rows[0].cells
    sc_hdr[0].text = 'Color'
    sc_hdr[1].text = 'Meaning'
    for cell in sc_hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    
    sc_data = [
        ('Green', 'Active, Completed, Paid'),
        ('Red', 'Inactive, Cancelled, Overdue'),
        ('Yellow', 'Pending, Warning'),
        ('Blue', 'In Progress, Info'),
        ('Gray', 'Draft, Neutral')
    ]
    
    for i, (color, meaning) in enumerate(sc_data, 1):
        row = sc_table.rows[i].cells
        row[0].text = color
        row[1].text = meaning
    
    doc.add_page_break()
    
    # Conclusion
    doc.add_heading('Conclusion', 1)
    p = doc.add_paragraph()
    p.add_run('Alafia Point Wellness Clinic Management System').bold = True
    p.add_run(' is designed to streamline your healthcare operations, improve patient care, and optimize clinic efficiency. With its comprehensive feature set, intuitive interface, and robust security, the system provides everything needed to manage a modern healthcare facility.')
    
    p = doc.add_paragraph()
    p.add_run('For additional assistance or feature requests, please contact our support team.')
    
    # Footer info
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('_' * 80)
    
    p = doc.add_paragraph()
    p.add_run('Document Version: ').bold = True
    p.add_run('1.0')
    
    p = doc.add_paragraph()
    p.add_run('Last Updated: ').bold = True
    p.add_run('March 2026')
    
    p = doc.add_paragraph()
    p.add_run('System Version: ').bold = True
    p.add_run('Alafia Point Wellness Clinic Management System')
    
    p = doc.add_paragraph()
    p.add_run('© 2026 Alafia Point Wellness Clinic. All rights reserved.')
    
    # Save the document
    doc.save('Alafia_Point_System_Documentation.docx')
    print("Word document created successfully: Alafia_Point_System_Documentation.docx")

if __name__ == '__main__':
    create_word_document()
